
# ------ WARNING!!! RUN THIS JUST ONCE!!! 
#import nltk
#nltk.download()

import os
import re
import pandas as pd
from docx import Document
import nltk
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

# ------------------- CONFIG -------------------
quarter_start = 1
year_start = 2021
quarter_end = 3
year_end = 2025

FINANCIAL_TERMS = [
    "revenue", "net income", "ebitda", "cet1", "capital ratio", "risk-weighted assets",
    "cash flow", "provision for credit losses", "loan growth", "net interest income",
    "net margin", "liquidity", "charge-off", "allowance for loan losses"
]

REGULATORY_THEMES = {
    "Capital Adequacy & Buffers": [
        "common equity tier 1", "cet1", "capital ratio", "risk-weighted assets", "tier 1 capital"
    ],
    "Liquidity": [
        "liquidity coverage ratio", "lcr", "net stable funding ratio", "nsfr", "liquidity buffer"
    ],
    "Credit Risk & Loan Quality": [
        "non-performing loans", "npl", "charge-off", "loan loss reserves", "reserve build"
    ],
    "Stress Testing & Scenario Planning": [
        "ccar", "dfast", "stress test", "scenario planning"
    ],
    "Supervisory & Regulatory References": [
        "regulatory requirement", "supervisory expectations", "living will", "gsib", "sifi"
    ],
    "Regulators": [
        "federal reserve", "occ", "fdic", "ecb", "eba", "boe"
    ],
    "Frameworks": [
        "basel iii", "basel iv", "dodd-frank"
    ],
    "Internal Metrics": [
        "cet1 ratio", "lcr", "nsfr", "leverage ratio", "net interest margin"
    ]
}

# ----------------- UTILITIES -----------------

def extract_citi_speakers(doc):
    citi_speakers = set()
    text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    # Flexible host match
    host_match = re.search(
        r"Host\s+([A-Z][a-z]+\s+[A-Z][a-z]+),\s*(?:Citi\s*)?(?:Head\s*(?:of)?\s*(?:Citi\s*)?)?\s*Investor\s*Relations",
        text, re.IGNORECASE)
    if host_match:
        citi_speakers.add(host_match.group(1).strip())

    # Extract speakers
    speaker_section = re.search(
        r"Speakers\s+(.*?)(?:PRESENTATION|\n[A-Z]{2,}:)",
        text, re.DOTALL | re.IGNORECASE)
    if speaker_section:
        speakers_text = speaker_section.group(1)
        for line in speakers_text.split("\n"):
            match = re.match(r"([A-Z][a-z]+\s+[A-Z][a-z]+),\s*Citi", line.strip())
            if match:
                citi_speakers.add(match.group(1).strip())

    return {s.lower() for s in citi_speakers}


def extract_quarter_year(filename):
    match = re.match(r"Q(\d)(\d{4})", filename)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None, None

def is_within_range(q, y):
    return (y > year_start or (y == year_start and q >= quarter_start)) and \
           (y < year_end or (y == year_end and q <= quarter_end))

def is_noise_line(text):
    lower = text.lower().strip()
    if re.fullmatch(r"t\s*r\s*a\s*n\s*s\s*c\s*r\s*i\s*p\s*t", lower.replace(" ", "")):
        return True
    if any(fragment in lower for fragment in [
        "transcript", "earnings review", "earnings call",
        "citi first quarter", "citi second quarter", "citi third quarter", "citi fourth quarter"
    ]):
        return True
    if "copyright" in lower and re.search(r"\d", lower):
        return True
    if lower.isdigit() and len(lower) < 4:
        return True
    if re.search(r"\b(january|february|march|april|may|june|july|august|"
                 r"september|october|november|december)\b.*\d{4}", lower):
        return True
    return False

def extract_financial_phrases(text):
    monetary_pattern = r"\$[\d,.]+(?:\s?(million|billion|trillion))?"
    monetary_matches = re.findall(monetary_pattern, text, re.IGNORECASE)
    term_matches = [term for term in FINANCIAL_TERMS if re.search(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE)]
    combined = list(set(term_matches + [m[0] if isinstance(m, tuple) else m for m in monetary_matches]))
    return ', '.join(combined)

def analyze_regulatory_themes(text):
    results = {}
    text_lower = text.lower()
    for theme, keywords in REGULATORY_THEMES.items():
        matches = [kw for kw in keywords if re.search(r'\b' + re.escape(kw) + r'\b', text_lower)]
        results[f"Mentions {theme}"] = "Yes" if matches else "No"
        results[f"{theme} phrase"] = ", ".join(matches)
    return results

def parse_transcript(doc_path, quarter, year):
    doc = Document(doc_path)
    citi_speakers = extract_citi_speakers(doc)
    raw_text = "\n".join(
        p.text.strip() for p in doc.paragraphs
        if p.text.strip() and not is_noise_line(p.text)
    )
    speaker_segments = re.findall(
        r"([A-Z][A-Z\s]+):\s*(.*?)(?=(?:[A-Z][A-Z\s]+:)|$)",
        raw_text, re.DOTALL
    )
    data = []
    for speaker, text in speaker_segments:
        clean_text = ' '.join(text.strip().split())
        speaker_name = speaker.title().strip()
        is_citi = "yes" if speaker_name.lower() in citi_speakers else "no"
        entry = {
            "year": year,
            "quarter": f"Q{quarter}",            
            "speaker": speaker_name,
            "text": clean_text,
            "financial_terms": extract_financial_phrases(clean_text),
            "is_citi_member": is_citi
        }
        entry.update(analyze_regulatory_themes(clean_text))
        data.append(entry)
    return data

USELESS_PHRASES = [
    "thanks", "thank you", "you're welcome", "sure", "okay", "yes", "no",
    "good morning", "good afternoon", "appreciate that", "great, thank you",
    "hi everyone", "hi", "hello", "bye", "have a good day", "Yeah.", 
    "All right.", "Hey.", "Yeah, sure.", "We get it.", "Enjoy the day.","Bye"
]

def is_useless_sentence(sentence):
    stripped = sentence.strip().lower()

    # Remove short, polite, or transitional phrases
    if len(stripped.split()) <= 5 and any(stripped.startswith(p) for p in USELESS_PHRASES):
        return True

    # Remove number-only or year-only lines
    if re.fullmatch(r"[^\w]*(\d{1,4})([^\w]*)", stripped) or stripped.isdigit():
        return True

    return False


# ---------------- MAIN PROCESSING ----------------

all_transcripts = []

for filename in os.listdir():
    if filename.endswith(".docx") and re.match(r"Q\d{5}", filename):
        quarter, year = extract_quarter_year(filename)
        if quarter and year and is_within_range(quarter, year):
            print(f"Processing: {filename}")
            all_transcripts.extend(parse_transcript(filename, quarter, year))

df_all = pd.DataFrame(all_transcripts)

# -------------- SPLIT INTO SENTENCES --------------
sentence_rows = []
for row in df_all.itertuples(index=False):
    sentences = sent_tokenize(row.text)
    for sentence in sentences:
        if is_useless_sentence(sentence):
            continue 
        entry = {
            "year": row.year,
            "quarter": row.quarter,
            "speaker": row.speaker,
            "sentence": sentence.strip(),            
            "financial_terms": extract_financial_phrases(sentence),
            "is_citi_member": row.is_citi_member
        }
        entry.update(analyze_regulatory_themes(sentence))
        sentence_rows.append(entry)

df_sentences = pd.DataFrame(sentence_rows)


# ---------------- EXPORT TO CSV ----------------
df_all.to_csv("transcript_full_phrases.csv", index=False)
df_sentences.to_csv("transcript_sentences.csv", index=False)



