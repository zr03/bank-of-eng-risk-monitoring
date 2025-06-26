# -*- coding: utf-8 -*-
"""
Created on Thu May 29 18:51:20 2025

@author: astri
"""

#pip install python-docx pytesseract pillow

import os
import re
from docx import Document
from docx.shared import Pt
from PIL import Image
from io import BytesIO
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def ocr_images_in_docx(input_path, output_path):
    doc = Document(input_path)
    rels = doc.part._rels

    for rel_id in list(rels):
        rel = rels[rel_id]
        if "image" in rel.target_ref:
            image_blob = rel.target_part.blob
            image = Image.open(BytesIO(image_blob))

            # Run OCR
            ocr_text = pytesseract.image_to_string(image).strip()
            if not ocr_text:
                ocr_text = "[IMAGE UNREADABLE]"

            # Replace image with OCR text
            for para in doc.paragraphs:
                if para._element.xpath(".//a:blip"):
                    para.clear()
                    run = para.add_run(ocr_text)
                    run.font.size = Pt(10)
                    break  # Replace only first match per image

    doc.save(output_path)
    print(f"✅ Saved OCR-rewritten file: {output_path}")

# -------- Run on Known Files --------
files = [f for f in os.listdir() if re.match(r"Q\d{5}\.docx", f)]

for filename in files:
    name, ext = os.path.splitext(filename)
    output_file = f"{name}_ocr{ext}"
    ocr_images_in_docx(filename, output_file)
